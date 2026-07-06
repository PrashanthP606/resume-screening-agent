import os
import re
from pypdf import PdfReader

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def clean_text(text: str) -> str:
    """Clean extracted text by removing excessive whitespace and double line breaks."""
    if not text:
        return ""
    # Strip each line to prevent spaces before/after newlines from being collapsed into word spacing
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(lines)
    # Collapse multiple spaces
    text = re.sub(r'[ \t]+', ' ', text)
    # Collapse multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pypdf."""
    text_content = []
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    return clean_text("\n".join(text_content))

def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx with a fallback."""
    if not HAS_DOCX:
        raise ImportError(
            "python-docx is not installed. Please install it using 'pip install python-docx'."
        )
    
    text_content = []
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_content.append(paragraph.text)
        
        # Also parse tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text_content.append(" | ".join(row_text))
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
        
    return clean_text("\n".join(text_content))

def parse_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return clean_text(f.read())
    except Exception as e:
        raise ValueError(f"Failed to parse plain text file: {str(e)}")

def parse_resume(file_path: str) -> str:
    """Parse resume and extract raw text based on file extension."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        # Standardize docx parsing
        return parse_docx(file_path)
    elif ext in [".txt", ".md", ".rtf"]:
        return parse_txt(file_path)
    else:
        # Fallback to plain text parsing if it is a text-like format, otherwise error
        try:
            return parse_txt(file_path)
        except Exception:
            raise ValueError(f"Unsupported file format: {ext}")
